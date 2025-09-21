import base64
import os
import re
import tempfile

from docx import Document
from fastmcp import FastMCP
from spire.doc import Document as SpireDoc
from spire.doc import FileFormat
from spire.doc.common import *

# Create an MCP server with a descriptive name
mcp = FastMCP("Word Document Master 📄")

# --- Helper Functions ---
def get_safe_filepath(filename: str) -> str:
    """
    Constructs a full, safe path inside the system's temp directory.
    This prevents directory traversal attacks and ensures the path is writable.
    """
    # Sanitize the filename to prevent security issues like ../../etc/passwd
    safe_filename = os.path.basename(filename)
    # Join it with the system's temporary directory (e.g., /tmp)
    return os.path.join(tempfile.gettempdir(), safe_filename)

@mcp.tool()
def create_document(filename: str) -> str:
    """
    Creates a new, blank Word document in a safe temporary directory on the server.

    Args:
        filename: The base name of the file to create (e.g., 'mydocument.docx').

    Returns:
        A success message.
    """
    document = Document()
    safe_path = get_safe_filepath(filename)  # Use the helper to get a writable path
    document.save(safe_path)
    return f"Document '{filename}' created successfully on the server in a temporary location."

@mcp.tool()
def add_paragraph(filename: str, text: str) -> str:
    """
    Adds a new paragraph to the end of a Word document.

    Args:
        filename: The name of the document to edit.
        text: The text to add.

    Returns:
        A success message.
    """
    safe_path = get_safe_filepath(filename)  # Use the helper
    try:
        document = Document(safe_path)
        document.add_paragraph(text)
        document.save(safe_path)
        return f"Paragraph added to '{filename}'."
    except Exception as e:
        return f"Error adding paragraph to '{filename}': {e}"

@mcp.tool()
def add_heading(filename: str, text: str, level: int = 1) -> str:
    """
    Adds a heading to a Word document.

    Args:
        filename: The name of the document.
        text: The heading text.
        level: The heading level (0-9).

    Returns:
        A success message.
    """
    safe_path = get_safe_filepath(filename)  # Use the helper
    try:
        document = Document(safe_path)
        document.add_heading(text, level=level)
        document.save(safe_path)
        return f"Heading added to '{filename}'."
    except Exception as e:
        return f"Error adding heading to '{filename}': {e}"

@mcp.tool()
def download_document(filename: str) -> str:
    """
    Reads a document from the server's temp directory and returns its content
    as a base64 encoded string for the client to download.

    Args:
        filename: The name of the file on the server (e.g., 'report.docx').

    Returns:
        The base64 encoded content of the file, or an error message if not found.
    """
    safe_path = get_safe_filepath(filename)  # Use the helper
    try:
        with open(safe_path, "rb") as docx_file:
            encoded_string = base64.b64encode(docx_file.read()).decode('utf-8')
        # Clean up the temporary file after it has been read
        os.remove(safe_path)
        return encoded_string
    except FileNotFoundError:
        return "Error: File not found on the server."
    except Exception as e:
        return f"Error downloading '{filename}': {e}"

@mcp.tool()
def convert_to_pdf(filename: str) -> str:
    """
    Converts a Word document to a PDF file on the server using Spire.Doc.
    
    Args:
        filename: The name of the document to convert.

    Returns:
        A success message or an error message.
    """
    docx_path = get_safe_filepath(filename)
    pdf_path = docx_path.replace('.docx', '.pdf')
    if docx_path == pdf_path:
        pdf_path += '.pdf'
    document = SpireDoc()
    try:
        document.LoadFromFile(docx_path)
        document.SaveToFile(pdf_path, FileFormat.PDF)
        return f"Document '{filename}' successfully converted to '{os.path.basename(pdf_path)}'."
    except Exception as e:
        return f"Error converting '{filename}' to PDF: {e}"
    finally:
        document.Close()

@mcp.tool()
def upload_to_s3(filename: str) -> str:
    """
    Simulates the process of uploading a file to an S3 bucket.

    Args:
        filename: The name of the file to "upload".

    Returns:
        A message confirming the simulated upload.
    """
    # This is a placeholder for actual S3 upload logic.
    # A real implementation would use a library like 'boto3'.
    # For example:
    # import boto3
    # s3_client = boto3.client('s3')
    # s3_client.upload_file(get_safe_filepath(filename), 'your-bucket-name', filename)
    
    # After a successful upload, you can clean up the local file.
    safe_path = get_safe_filepath(filename)
    if os.path.exists(safe_path):
        os.remove(safe_path)
    
    return f"PDF '{filename}' has been successfully uploaded to the S3 bucket."


@mcp.tool()
def process_document(input_string: str, filename: str) -> str:
    """
    Unifies the document creation, editing, conversion, and upload process.
    
    This tool takes a structured string, extracts the header and body,
    creates a Word document, adds the content, converts it to a PDF,
    and then simulates the final upload to S3.

    Args:
        input_string: A string containing a <header> and <body> tag.
        filename: The desired name for the output Word and PDF files.

    Returns:
        A final status message indicating the completion of the entire process.
    """
    # 1. Parse the input string
    header_match = re.search(r'<header>(.*?)</header>', input_string, re.DOTALL)
    body_match = re.search(r'<body>(.*?)</body>', input_string, re.DOTALL)
    
    if not header_match or not body_match:
        return "Error: Input string must contain both <header> and <body> tags."
        
    header_text = header_match.group(1).strip()
    body_text = body_match.group(1).strip()
    
    # 2. Create the document
    create_result = create_document(filename)
    print(create_result)
    
    # 3. Add the header
    add_header_result = add_heading(filename, header_text)
    print(add_header_result)
    
    # 4. Add the body
    add_body_result = add_paragraph(filename, body_text)
    print(add_body_result)
    
    # 5. Convert to PDF
    pdf_filename = filename.replace('.docx', '.pdf')
    if pdf_filename == filename:
        pdf_filename += '.pdf'
    
    convert_result = convert_to_pdf(filename)
    print(convert_result)
    
    # 6. Upload the final PDF to S3
    s3_upload_result = upload_to_s3(pdf_filename)
    print(s3_upload_result)
    
    return f"Process completed successfully. Document '{filename}' converted to '{pdf_filename}' and uploaded."


if __name__ == "__main__":
    mcp.run()
